#!/usr/bin/env python3
"""Convert a markdown draft to .docx with proper formatting and Track Changes enabled.

The output .docx opens in Word with Track Changes already on — ready for further editing.
[[wikilinks]] are resolved to (Author, Year) citations using sources/{stem}.md frontmatter.
Unresolvable wikilinks are kept as [stem] italic placeholders.

Usage:
    python3 scripts/convert_to_docx.py projects/{slug}/Drafts/{section}-v1.draft.md
    python3 scripts/convert_to_docx.py path/to/draft.md --output path/to/output.docx
    python3 scripts/convert_to_docx.py path/to/draft.md --no-track-changes
    python3 scripts/convert_to_docx.py path/to/draft.md --no-resolve-citations
"""

from __future__ import annotations

import re
import sys
from argparse import ArgumentParser
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]


def _resolve_citations_in_text(text: str) -> str:
    """Resolve [[wikilinks]] to (Author, Year) using resolve_citations module."""
    try:
        sys.path.insert(0, str(ROOT / "scripts"))
        from resolve_citations import resolve_all  # type: ignore
        resolved_text, _, _ = resolve_all(text)
        return resolved_text
    except ImportError:
        return text


# ---------------------------------------------------------------------------
# Markdown parser (minimal — handles the patterns used in drafts)
# ---------------------------------------------------------------------------

def parse_markdown(text: str) -> list[dict]:
    """Parse markdown into a flat list of block nodes.

    Node types: heading, paragraph, listitem, hr, blank
    """
    blocks = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            blocks.append({"type": "blank"})
            i += 1
            continue

        # Horizontal rule / page break
        if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", stripped):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # ATX headings: # ## ### ####
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            content = m.group(2).strip()
            blocks.append({"type": "heading", "level": level, "text": content})
            i += 1
            continue

        # Setext headings (=== or ---)
        if i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if re.match(r"^=+$", next_line):
                blocks.append({"type": "heading", "level": 1, "text": stripped})
                i += 2
                continue
            if re.match(r"^-+$", next_line) and len(next_line) > 2:
                blocks.append({"type": "heading", "level": 2, "text": stripped})
                i += 2
                continue

        # Unordered list item
        m = re.match(r"^[-*+]\s+(.*)", stripped)
        if m:
            blocks.append({"type": "listitem", "ordered": False, "text": m.group(1), "indent": 0})
            i += 1
            continue

        # Ordered list item
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            blocks.append({"type": "listitem", "ordered": True, "text": m.group(1), "indent": 0})
            i += 1
            continue

        # Block quote
        m = re.match(r"^>\s*(.*)", stripped)
        if m:
            blocks.append({"type": "blockquote", "text": m.group(1)})
            i += 1
            continue

        # Regular paragraph — may span multiple lines
        para_lines = [stripped]
        while i + 1 < len(lines):
            next_stripped = lines[i + 1].strip()
            if (not next_stripped or
                    re.match(r"^#{1,4}\s", next_stripped) or
                    re.match(r"^[-*+]\s", next_stripped) or
                    re.match(r"^\d+\.\s", next_stripped) or
                    re.match(r"^-{3,}$", next_stripped)):
                break
            i += 1
            para_lines.append(next_stripped)
        blocks.append({"type": "paragraph", "text": " ".join(para_lines)})
        i += 1

    return blocks


def parse_inline(text: str) -> list[dict]:
    """Parse inline markdown into runs: plain, bold, italic, bold-italic, wikilink, code."""
    runs = []
    # Pattern order matters: bold-italic first, then bold, then italic
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"      # bold-italic
        r"|(\*\*(.+?)\*\*)"          # bold
        r"|(__(.+?)__)"              # bold alt
        r"|(\*(.+?)\*)"             # italic
        r"|(_(.+?)_)"               # italic alt
        r"|(`(.+?)`)"               # code
        r"|(\[\[(.+?)\]\])"         # wikilink
    )
    last = 0
    for m in pattern.finditer(text):
        # Plain text before this match
        if m.start() > last:
            runs.append({"type": "plain", "text": text[last:m.start()]})

        if m.group(1):    # bold-italic
            runs.append({"type": "bold-italic", "text": m.group(2)})
        elif m.group(3):  # **bold**
            runs.append({"type": "bold", "text": m.group(4)})
        elif m.group(5):  # __bold__
            runs.append({"type": "bold", "text": m.group(6)})
        elif m.group(7):  # *italic*
            runs.append({"type": "italic", "text": m.group(8)})
        elif m.group(9):  # _italic_
            runs.append({"type": "italic", "text": m.group(10)})
        elif m.group(11): # `code`
            runs.append({"type": "code", "text": m.group(12)})
        elif m.group(13): # [[wikilink]]
            runs.append({"type": "wikilink", "text": m.group(14)})

        last = m.end()

    if last < len(text):
        runs.append({"type": "plain", "text": text[last:]})

    return runs if runs else [{"type": "plain", "text": text}]


# ---------------------------------------------------------------------------
# python-docx writer
# ---------------------------------------------------------------------------

def enable_track_changes(doc) -> None:
    """Add w:trackChanges to document settings so Word opens with Track Changes on."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    settings_element = doc.settings.element
    # Only add if not already present
    existing = settings_element.find(qn("w:trackChanges"))
    if existing is None:
        track = OxmlElement("w:trackChanges")
        settings_element.insert(0, track)


def set_run_style(run, run_type: str) -> None:
    """Apply formatting to a docx run based on parsed inline type."""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    if run_type in ("bold", "bold-italic"):
        run.bold = True
    if run_type in ("italic", "bold-italic"):
        run.italic = True
    if run_type == "code":
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    if run_type == "wikilink":
        # Unresolved wikilink — style as italicised placeholder [stem]
        run.italic = True
        run.text = f"[{run.text}]"


def add_paragraph_with_inline(doc_or_cell, text: str, style: str = "Normal"):
    """Add a paragraph with inline formatting applied."""
    para = doc_or_cell.add_paragraph(style=style)
    para.paragraph_format.space_after = None
    for run_data in parse_inline(text):
        run = para.add_run(run_data["text"])
        set_run_style(run, run_data["type"])
    return para


def blocks_to_docx(blocks: list[dict], doc) -> None:
    """Write parsed markdown blocks into the docx document."""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    heading_styles = {
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3",
        4: "Heading 4",
    }

    for block in blocks:
        btype = block["type"]

        if btype == "blank":
            continue  # Skip blank lines (spacing handled by paragraph spacing)

        elif btype == "heading":
            level = block["level"]
            style = heading_styles.get(level, "Heading 4")
            para = doc.add_heading(block["text"], level=level)

        elif btype == "paragraph":
            add_paragraph_with_inline(doc, block["text"], style="Normal")

        elif btype == "blockquote":
            # Style as indented italic paragraph
            para = doc.add_paragraph(style="Normal")
            para.paragraph_format.left_indent = Inches(0.5)
            run = para.add_run(block["text"])
            run.italic = True

        elif btype == "listitem":
            style = "List Bullet" if not block.get("ordered") else "List Number"
            add_paragraph_with_inline(doc, block["text"], style=style)

        elif btype == "hr":
            # Add a thin paragraph with a bottom border as separator
            para = doc.add_paragraph()
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)


def add_document_header(doc, source_path: Path, track_changes: bool) -> None:
    """Add a comment header to the document explaining its origin and Track Changes status."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    para = doc.add_paragraph()
    run = para.add_run(
        f"Research Knowledge System — Auto-converted from: {source_path.name}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    if track_changes:
        note_para = doc.add_paragraph()
        run2 = note_para.add_run(
            "⚠ Track Changes is enabled in this document. "
            "All edits will be recorded. "
            "Accept/reject changes via Review → Accept All / Reject All."
        )
        run2.font.size = Pt(9)
        run2.bold = True
        run2.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# Main conversion function
# ---------------------------------------------------------------------------

def convert_md_to_docx(
    input_path: Path,
    output_path: Path | None = None,
    track_changes: bool = True,
    resolve_citations: bool = True,
    verbose: bool = True,
) -> Path:
    """Convert a markdown file to .docx. Returns the output path."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        sys.exit("python-docx is required: pip install python-docx")

    if not input_path.exists():
        sys.exit(f"Input file not found: {input_path}")

    if output_path is None:
        output_path = input_path.with_suffix(".docx")

    # Read and parse
    text = input_path.read_text(encoding="utf-8")

    # Strip YAML frontmatter if present
    if text.startswith("---"):
        end = text.find("\n---", 4)
        if end != -1:
            text = text[end + 4:].lstrip("\n")

    # Resolve [[wikilinks]] to (Author, Year) before converting
    if resolve_citations:
        resolved_text = _resolve_citations_in_text(text)
        if resolved_text != text and verbose:
            import re as _re
            before = len(_re.findall(r"\[\[.+?\]\]", text))
            after = len(_re.findall(r"\[\[.+?\]\]", resolved_text))
            print(f"  Citations: {before - after} resolved, {after} unresolved")
        text = resolved_text

    blocks = parse_markdown(text)

    # Build document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Set margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Header
    add_document_header(doc, input_path, track_changes)

    # Content
    blocks_to_docx(blocks, doc)

    # Enable Track Changes
    if track_changes:
        enable_track_changes(doc)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    if verbose:
        rel = output_path.relative_to(ROOT) if ROOT in output_path.parents else output_path
        print(f"✓ Saved: {rel}")
        if track_changes:
            print("  Track Changes: ON — edits in Word will be recorded automatically.")

    return output_path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    parser = ArgumentParser(
        description="Convert a markdown draft to .docx with optional Track Changes enabled.",
    )
    parser.add_argument("input", metavar="DRAFT.md",
                        help="Path to the markdown draft file.")
    parser.add_argument("--output", metavar="OUTPUT.docx",
                        help="Output .docx path. Default: same directory as input.")
    parser.add_argument("--no-track-changes", action="store_true", dest="no_track",
                        help="Disable Track Changes in the output document.")
    parser.add_argument("--no-resolve-citations", action="store_true", dest="no_resolve",
                        help="Keep [[wikilinks]] as-is instead of resolving to (Author, Year).")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    output_path = Path(args.output).resolve() if args.output else None
    track_changes = not args.no_track
    resolve_citations = not args.no_resolve

    out = convert_md_to_docx(
        input_path, output_path,
        track_changes=track_changes,
        resolve_citations=resolve_citations,
    )
    print(f"\nOpen in Word and edit with Track Changes {'already enabled' if track_changes else 'disabled'}.")
    print(f"Path: {out}")


if __name__ == "__main__":
    main()
